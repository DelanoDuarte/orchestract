import io
import os
from functools import lru_cache

from docx import Document as DocxDocument
from google import genai
from google.genai import types
from google.genai.types import HttpOptions

# Default Gemini model used for document/contract summarization. Kept as a
# module constant (mirroring AI_MODEL for the Anthropic assistant) so the model
# string can't be smuggled in from user input; overridable via env for ops.
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")

# Word .docx. Gemini can't read this inline (it's a zip of XML, not a media
# type the model ingests), so we extract its text locally first -- see
# build_document_part / extract_docx_text below.
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# MIME types we can hand straight to Gemini as an inline file part. Gemini is
# natively multimodal, so real uploaded contracts (PDFs, scans) can be
# summarized without a local text-extraction step -- a big step up from the
# text-only summarizer this replaces.
_INLINE_MIME_TYPES = frozenset(
    {
        "text/plain",
        "text/markdown",
        "application/pdf",
        "image/png",
        "image/jpeg",
        "image/webp",
    }
)

# Everything we can summarize: types Gemini reads inline, plus .docx which we
# convert to text ourselves. The application layer gates on this set.
SUMMARIZABLE_MIME_TYPES = _INLINE_MIME_TYPES | {DOCX_MIME}


def extract_docx_text(data: bytes) -> str:
    """Pull the readable text out of a .docx: paragraphs in order, plus table
    rows flattened to pipe-separated cells (contracts often keep pricing and
    party details in tables). Returns '' for an empty document."""
    document = DocxDocument(io.BytesIO(data))
    lines = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


def build_document_part(data: bytes, content_type: str, name: str) -> types.Part | str:
    """Turn raw document bytes into a piece of model input: an inline Part for
    natively-supported media, or extracted plain text for .docx. Returns a
    string so callers can drop it straight into a `contents` list."""
    if content_type == DOCX_MIME:
        text = extract_docx_text(data)
        return f"Contents of '{name}':\n{text}" if text else f"'{name}' is an empty document."
    return types.Part.from_bytes(data=data, mime_type=content_type)


def _use_vertexai() -> bool:
    return os.environ.get("GOOGLE_GENAI_USE_VERTEXAI", "").strip().lower() in {"1", "true", "yes"}


def _api_key() -> str | None:
    # Gemini Developer API (AI Studio) key. GOOGLE_API_KEY is the SDK's own
    # convention; GEMINI_API_KEY is accepted as a common alias.
    return os.environ.get("GOOGLE_API_KEY") or os.environ.get("GEMINI_API_KEY")


def gemini_enabled() -> bool:
    """True when summarization can reach a Gemini backend. Two options:

    * Developer API (AI Studio): a single GOOGLE_API_KEY / GEMINI_API_KEY --
      the simplest to run on non-GCP hosts (e.g. Railway), no service account.
    * Vertex AI via ADC: the Vertex opt-in plus a project. We don't inspect ADC
      credentials here (google-auth resolves those lazily at call time), so a
      half-configured Vertex environment still surfaces a clear "not configured"
      error rather than a cryptic auth failure deep in the SDK.
    """
    if _api_key():
        return True
    return _use_vertexai() and bool(os.environ.get("GOOGLE_CLOUD_PROJECT"))


@lru_cache
def get_gemini_client() -> genai.Client:
    # Vertex opt-in wins when set: the SDK reads GOOGLE_CLOUD_PROJECT /
    # GOOGLE_CLOUD_LOCATION and authenticates via Application Default
    # Credentials, so there's no key in app config to log or leak.
    if _use_vertexai():
        return genai.Client(http_options=HttpOptions(api_version="v1"), vertexai=True)
    # Otherwise use the Developer API with an explicit key (no ADC, no service
    # account) -- passed in rather than left to the SDK's env lookup so the
    # GEMINI_API_KEY alias works too.
    return genai.Client(api_key=_api_key())


async def generate_summary(contents: list[types.Part | str], system_instruction: str) -> str:
    """One-shot summarization under a fixed system prompt. `contents` interleaves
    inline file parts (the actual uploaded documents) with text instructions.
    Async so it doesn't block the event loop (uses the SDK's `aio` surface)."""
    client = get_gemini_client()
    response = await client.aio.models.generate_content(
        model=GEMINI_MODEL,
        contents=contents,
        config=types.GenerateContentConfig(system_instruction=system_instruction),
    )
    return (response.text or "").strip()
