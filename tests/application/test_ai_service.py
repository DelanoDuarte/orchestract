import asyncio

import pytest

from app.application.ai_service import AIService
from app.domain.ai.exceptions import AIUnavailableError, NoVersionsError, UnsupportedContentTypeError
from app.domain.contracts.models import Contract
from app.domain.documents.models import Document
from app.domain.tenancy.models import Organization
from app.domain.tenancy.plans import Plan


class _FakeDocumentService:
    def __init__(self, document: Document) -> None:
        self._document = document

    async def get(self, document_id: int) -> Document:
        return self._document


class _FakeContractService:
    def __init__(self, contract: Contract) -> None:
        self._contract = contract

    async def get(self, contract_id: int) -> Contract:
        return self._contract


def _contract(ai_config: dict | None = None) -> Contract:
    contract = Contract.create(1, "MSA", "Services")
    if ai_config is not None:
        contract.ai_config = ai_config
    return contract


def _service_for(document: Document, contract: Contract | None = None) -> AIService:
    return AIService(
        contract_service=_FakeContractService(contract or _contract()),
        document_service=_FakeDocumentService(document),
        role_service=None,
        workflow_service=None,
    )


def _org(plan: Plan = Plan.BUSINESS) -> Organization:
    return Organization(plan=plan.value)


def _enable_gemini(monkeypatch):
    """Configure the Vertex-AI-via-ADC env that gemini_enabled() checks."""
    monkeypatch.setenv("GOOGLE_GENAI_USE_VERTEXAI", "true")
    monkeypatch.setenv("GOOGLE_CLOUD_PROJECT", "test-project")


def test_summarize_document_requires_a_configured_provider(monkeypatch):
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    monkeypatch.delenv("GOOGLE_GENAI_USE_VERTEXAI", raising=False)
    monkeypatch.delenv("GOOGLE_CLOUD_PROJECT", raising=False)
    document = Document.create(name="MSA")
    service = _service_for(document)

    with pytest.raises(AIUnavailableError):
        asyncio.run(service.summarize_document(1, _org()))


def test_summarize_document_requires_gemini_configured(monkeypatch):
    # Anthropic (assistant) alone doesn't enable summaries -- those need Gemini.
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")
    monkeypatch.delenv("GOOGLE_GENAI_USE_VERTEXAI", raising=False)
    monkeypatch.delenv("GOOGLE_CLOUD_PROJECT", raising=False)
    document = Document.create(name="MSA")
    service = _service_for(document)

    with pytest.raises(AIUnavailableError):
        asyncio.run(service.summarize_document(1, _org()))


def test_summarize_document_requires_a_plan_with_ai_enabled(monkeypatch):
    _enable_gemini(monkeypatch)
    document = Document.create(name="MSA")
    service = _service_for(document)

    with pytest.raises(AIUnavailableError):
        asyncio.run(service.summarize_document(1, _org(Plan.FREE)))


def test_summarize_document_rejects_a_document_with_no_versions(monkeypatch):
    _enable_gemini(monkeypatch)
    document = Document.create(name="MSA")
    service = _service_for(document)

    with pytest.raises(NoVersionsError):
        asyncio.run(service.summarize_document(1, _org()))


def test_summarize_document_rejects_unsupported_content_types(monkeypatch):
    _enable_gemini(monkeypatch)
    document = Document.create(name="MSA")
    document.add_version(
        storage_connection_id=1,
        storage_key="org-1/contract-1/doc-1/v1-msa.zip",
        original_filename="msa.zip",
        content_type="application/zip",
        size_bytes=1024,
        uploaded_by="alice",
    )
    service = _service_for(document)

    with pytest.raises(UnsupportedContentTypeError):
        asyncio.run(service.summarize_document(1, _org()))


def test_is_available_org_plan_gate_beats_contract_level_force_on(monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")
    service = _service_for(Document.create(name="MSA"))
    contract = _contract(ai_config={"enabled": True})

    assert not service.is_available(_org(Plan.FREE), contract)


def test_is_available_contract_level_force_off_beats_allowing_org_plan(monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")
    service = _service_for(Document.create(name="MSA"))
    contract = _contract(ai_config={"enabled": False})

    assert not service.is_available(_org(Plan.BUSINESS), contract)


def test_is_available_defaults_to_org_plan_when_config_absent(monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")
    service = _service_for(Document.create(name="MSA"))

    assert service.is_available(_org(Plan.BUSINESS), _contract())
    assert not service.is_available(_org(Plan.FREE), _contract())


def test_summarize_document_sends_the_uploaded_file_to_gemini(monkeypatch):
    _enable_gemini(monkeypatch)
    document = Document.create(name="MSA")
    document.add_version(
        storage_connection_id=1,
        storage_key="org-1/contract-1/doc-1/v1-msa.pdf",
        original_filename="msa.pdf",
        content_type="application/pdf",
        size_bytes=13,
        uploaded_by="alice",
    )
    contract = _contract(ai_config={"instructions": "Flag any auto-renewal."})
    service = _service_for(document, contract)

    captured = {}

    async def _fake_generate_summary(contents, system_instruction):
        captured["contents"] = contents
        captured["system_instruction"] = system_instruction
        return "a summary"

    class _FakeDocumentServiceWithDownload(_FakeDocumentService):
        async def download_version(self, version):
            return b"%PDF-1.4 fake"

        async def set_summary(self, document_id, summary):
            captured["summary"] = summary
            return document

    service._document_service = _FakeDocumentServiceWithDownload(document)
    monkeypatch.setattr("app.application.ai_service.generate_summary", _fake_generate_summary)

    asyncio.run(service.summarize_document(1, _org()))

    part, prompt = captured["contents"]
    assert part.inline_data.mime_type == "application/pdf"
    assert part.inline_data.data == b"%PDF-1.4 fake"
    # Per-contract instructions ride along in the prompt.
    assert "Flag any auto-renewal." in prompt
    assert captured["summary"] == "a summary"


class _FakeToolRunner:
    def __init__(self, **kwargs) -> None:
        self.kwargs = kwargs

    def __aiter__(self):
        async def _empty():
            return
            yield  # pragma: no cover

        return _empty()


class _FakeBetaMessages:
    def __init__(self) -> None:
        self.last_kwargs: dict | None = None

    def tool_runner(self, **kwargs):
        self.last_kwargs = kwargs
        return _FakeToolRunner(**kwargs)


class _FakeBeta:
    def __init__(self) -> None:
        self.messages = _FakeBetaMessages()


class _FakeAssistantClient:
    def __init__(self) -> None:
        self.beta = _FakeBeta()


class _FakeUser:
    name = "Dana"


def test_run_assistant_restricts_tools_and_resolves_model_from_contract_config(monkeypatch):
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")
    contract = _contract(ai_config={"model": "claude-sonnet-5", "allowed_tools": []})
    contract_service = _FakeContractService(contract)
    service = AIService(
        contract_service=contract_service, document_service=None, role_service=None, workflow_service=None
    )
    fake_client = _FakeAssistantClient()
    monkeypatch.setattr("app.application.ai_service.get_anthropic_client", lambda: fake_client)

    asyncio.run(service.run_assistant(1, "do something", _FakeUser(), _org()))

    kwargs = fake_client.beta.messages.last_kwargs
    assert kwargs["model"] == "claude-sonnet-5"
    assert [t.name for t in kwargs["tools"]] == ["get_contract_state"]


def _docx_bytes() -> bytes:
    import io

    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Acme Corp and Damian LLC agree as follows.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Fee"
    table.rows[0].cells[1].text = "USD 10,000/mo"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_build_document_part_extracts_docx_text_instead_of_inlining():
    from app.infrastructure.ai.gemini import DOCX_MIME, SUMMARIZABLE_MIME_TYPES, build_document_part

    # .docx is summarizable but can't be sent to Gemini inline -- it's converted
    # to text, so the part is a str carrying the paragraphs and table cells.
    assert DOCX_MIME in SUMMARIZABLE_MIME_TYPES
    part = build_document_part(_docx_bytes(), DOCX_MIME, "MSA.docx")
    assert isinstance(part, str)
    assert "Acme Corp and Damian LLC" in part
    assert "Fee | USD 10,000/mo" in part


def test_build_document_part_inlines_native_media():
    from app.infrastructure.ai.gemini import build_document_part

    part = build_document_part(b"%PDF-1.4 fake", "application/pdf", "x.pdf")
    assert part.inline_data.mime_type == "application/pdf"
    assert part.inline_data.data == b"%PDF-1.4 fake"
